from docx import Document
import re
import os
import pandas as pd
from datetime import date
# ------------------- Helpers -------------------
DASH = "–"  # en-dash for year ranges

EXCEL_CELL_LIMIT = 32767  # Excel max char limit per cell

def split_into_excel_cells(text, limit=EXCEL_CELL_LIMIT):
    """Split text into chunks that fit within Excel cell limit."""
    if not text:
        return [""]
    return [text[i:i+limit] for i in range(0, len(text), limit)]

HEADER_LINE_RE = re.compile(
    r"""^\s*
        (?:[A-Za-z]\.)?
        (?:\d+(?:\.\d+)*)?
        [\.\)]?\s*
        (?:report\s*title|full\s*title|full\s*report\s*title|title\s*\(long[-\s]*form\))
        [\s:–-]*$
    """, re.I | re.X
)

def remove_emojis(text: str) -> str:
    """Universal emoji remover."""
    emoji_pattern = re.compile(
        "[" 
        "\U0001F600-\U0001F64F"  # emoticons
        "\U0001F300-\U0001F5FF"  # symbols & pictographs
        "\U0001F680-\U0001F6FF"  # transport & map
        "\U0001F700-\U0001F77F"  # alchemical
        "\U0001F780-\U0001F7FF"  # geometric
        "\U0001F800-\U0001F8FF"  # arrows
        "\U0001F900-\U0001F9FF"  # supplemental
        "\U0001FA00-\U0001FAFF"  # chess, symbols
        "\U00002600-\U000026FF"  # misc symbols
        "\U00002700-\U000027BF"  # dingbats
        "\U00002B00-\U00002BFF"  # arrows & symbols
        "\U0001F1E0-\U0001F1FF"  # flags
        "]+", flags=re.UNICODE
    )
    return emoji_pattern.sub(r'', text or "")
def _remove_emojis(text: str) -> str:
    emoji_pattern = re.compile(
        "[" "\U0001F600-\U0001F64F"
        "\U0001F300-\U0001F5FF"
        "\U0001F680-\U0001F6FF"
        "\U0001F1E0-\U0001F1FF"
        "\U00002702-\U000027B0"
        "\U000024C2-\U0001F251" "]+",
        flags=re.UNICODE
    )
    return emoji_pattern.sub("", text)


def _norm(s: str) -> str:
    s = remove_emojis(s or "")
    return re.sub(r"\s+", " ", s.strip())

def _inline_title(text: str) -> str:
    m = re.split(r"[:\-–]", text, maxsplit=1)
    if len(m) > 1:
        right = m[1].strip()
        if right and not HEADER_LINE_RE.match(right):
            return right
    return ""

def _year_range_present(text: str) -> bool:
    return bool(re.search(r"20\d{2}\s*[\-–]\s*20\d{2}", text))

def _ensure_filename_start_and_year(title: str, filename: str) -> str:
    if not title.lower().startswith(filename.lower()):
        title = f"{filename} {title}"
    if not _year_range_present(title):
        title = f"{title} {DASH}2024–2030"
    return _norm(title)
# ------------------- Convert Paragraph to HTML -------------------
def paragraph_to_html(para):
    """Convert a docx paragraph into HTML with basic formatting."""
    text = para.text.strip()
    if not text:
        return ""

    # Check if it's a list item
    if para.style.name.lower().startswith("list"):
        return f"<li>{text}</li>"
    
    text = _remove_emojis(text)

    # Headings
    if para.style.name.startswith("Heading"):
        level = para.style.name.replace("Heading", "").strip()
        level = int(level) if level.isdigit() else 2
        return f"<h{level}>{text}</h{level}>"

    # Normal paragraph
    return f"<p>{text}</p>"

# def paragraph_to_html(para):
#     text = remove_emojis(para.text.strip())
#     if not text:
#         return ""
#     if para.style.name.lower().startswith("list"):
#         return f"<li>{text}</li>"
#     if para.style.name.startswith("Heading"):
#         level = para.style.name.replace("Heading", "").strip()
#         level = int(level) if level.isdigit() else 2
#         return f"<h{level}>{text}</h{level}>"
#     return f"<p>{text}</p>"

# ------------------- Extract Title -------------------
def extract_title(docx_path: str) -> str:
    doc = Document(docx_path)
    filename = os.path.splitext(os.path.basename(docx_path))[0]
    filename_low = filename.lower()
    blocks = [(p, (p.text or "").strip()) for p in doc.paragraphs if (p.text or "").strip()]

    capture = False
    for _, text in blocks:
        text = remove_emojis(text)
        if capture:
            return _ensure_filename_start_and_year(text, filename)
        if HEADER_LINE_RE.match(text):
            inline = _inline_title(text)
            if inline:
                return _ensure_filename_start_and_year(inline, filename)
            capture = True
            continue

    for table in doc.tables:
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_text = (cell.text or "").strip().lower()
                if not cell_text:
                    continue
                if "report title" in cell_text or "full title" in cell_text or "full report title" in cell_text:
                    if c_idx + 1 < len(row.cells):
                        nxt = row.cells[c_idx+1].text.strip()
                        if nxt:
                            return _ensure_filename_start_and_year(nxt, filename)
                    if r_idx + 1 < len(table.rows):
                        nxt = row.rows[r_idx+1].cells[c_idx].text.strip()
                        if nxt:
                            return _ensure_filename_start_and_year(nxt, filename)

    for _, text in blocks:
        low = text.lower()
        if low.startswith("full report title") or low.startswith("full title"):
            inline = _inline_title(text)
            if inline:
                return _ensure_filename_start_and_year(inline, filename)
        if low.startswith(filename_low) and "forecast" in low:
            return _ensure_filename_start_and_year(text, filename)

    return "Title Not Available"

# ------------------- Extract Description -------------------
def extract_description(docx_path):
    doc = Document(docx_path)
    html_output = []
    capture, inside_list = False, False  

    target_headings = [
        "introduction and strategic context",
        "market segmentation and forecast scope",
        "market trends and innovation landscape",
        "competitive intelligence and benchmarking",
        "regional landscape and adoption outlook",
        "end-user dynamics and use case",
        "recent developments + opportunities & restraints",
    ]

    def clean_heading(text):
        text = remove_emojis(text.strip())
        text = re.sub(r'^[^\w]+', '', text)  
        text = re.sub(r'(?i)section\s*\d+[:\-]?\s*', '', text)  
        text = re.sub(r'^\d+[\.\-\)]\s*', '', text)  
        text = re.sub(r'\s+', ' ', text)  
        return text.lower().strip()

    def run_to_html(run):
        text = remove_emojis(run.text.strip())
        if not text:
            return ""
        if run.bold and run.italic:
            return f"<b><i>{text}</i></b>"
        elif run.bold:
            return f"<b>{text}</b>"
        elif run.italic:
            return f"<i>{text}</i>"
        return text

    for para in doc.paragraphs:
        text = remove_emojis(para.text.strip())
        if not text:
            continue

        cleaned = clean_heading(text)

        if not capture and any(h in cleaned for h in target_headings):
            capture = True  

        if capture and "report summary, faqs, and seo schema" in cleaned:
            break  

        if capture:
            content = "".join(run_to_html(run) for run in para.runs if run.text.strip())

            if any(h in cleaned for h in target_headings):
                html_output.append("<br>")
                matched = next(h for h in target_headings if h in cleaned)
                html_output.append(f"<h2>{matched.title()}</h2>")
                continue

            if "list" in para.style.name.lower():
                if not inside_list:
                    html_output.append("<ul>")
                    inside_list = True
                html_output.append(f"<li>{content}</li>")
                continue
            else:
                if inside_list:
                    html_output.append("</ul>")
                    inside_list = False

            html_output.append(f"<p>{content}</p>")

    if inside_list:
        html_output.append("</ul>")
    return "\n".join(html_output)


# ------------------- TOC Extraction -------------------
def extract_toc(docx_path):
    doc = Document(docx_path)
    html_output, inside_list, capture = [], False, False
    end_reached = False

    for para in doc.paragraphs:
        text = remove_emojis(para.text.strip())
        low = text.lower()

        if not capture and "table of contents" in low:
            capture = True
            continue

        if capture:
            if "list of figures" in low:
                html_part = paragraph_to_html(para)
                if html_part:
                    html_output.append(html_part)  
                end_reached = True
                continue  

            if end_reached:
                style = getattr(para.style, "name","").lower()
                if "heading" in style or re.match(r"^\d+[\.\)]\s", text):
                    break  

            html_part = paragraph_to_html(para)
            if html_part:
                if html_part.startswith("<li>"):
                    if not inside_list:
                        html_output.append("<ul>")
                        inside_list = True
                    html_output.append(html_part)
                else:
                    if inside_list:
                        html_output.append("</ul>")
                        inside_list = False
                    html_output.append(html_part)

    if inside_list:
        html_output.append("</ul>")
    return "".join(html_output).strip()

# ------------------- FAQ Schema -------------------
def extract_faq_schema(docx_path):
    doc = Document(docx_path)
    paras = [remove_emojis(p.text) for p in doc.paragraphs if p.text and p.text.strip()]

    capture, started, depth, buf = False, False, 0, []

    for text in paras:
        line = text.strip()
        low = line.lower()

        if not capture:
            if "faq schema" in low:
                capture = True
            continue

        if not started:
            if low in ("json", "copyedit", "json copy", "json copyedit"):
                continue
            if "{" in line:
                started = True
            else:
                continue

        buf.append(line)
        depth += line.count("{") - line.count("}")
        if depth <= 0 and started:
            break

    return "\n".join(buf).strip()

def extract_methodology_from_faqschema(docx_path):
    faq_schema_str = extract_faq_schema(docx_path)
    if not faq_schema_str:
        return ""   
    
    try:
        faq_data = json.loads(faq_schema_str)
    except json.JSONDecodeError:
        return ""   
    
    faqs = []
    q_count = 0
    for item in faq_data.get("mainEntity", []):
        q_count += 1
        question = remove_emojis(item.get("name", "").strip())
        answer = remove_emojis(item.get("acceptedAnswer", {}).get("text", "").strip())
        if question and answer:
            faqs.append(
                f"<p><strong>Q{q_count}: {html.escape(question)}</strong><br>"
                f"A{q_count}: {html.escape(answer)}</p>"
            )
    return "\n".join(faqs)


# ------------------- Report Coverage -------------------
def extract_report_coverage_table_with_style(docx_path):
    doc = Document(docx_path)

    for table in doc.tables:
        first_row_text = " ".join([c.text.strip().lower() for c in table.rows[0].cells])
        if "report attribute" in first_row_text or "report coverage table" in first_row_text:
            html = []
            html.append('<h2><strong>7.1. Report Coverage Table</strong></h2>')
            html.append('<table cellspacing="0" style="border-collapse:collapse; width:100%"><tbody>')

            for r_idx, row in enumerate(table.rows):
                html.append("<tr>")
                for c_idx, cell in enumerate(row.cells):
                    text = remove_emojis(cell.text.strip())

                    bg = "#deeaf6" if r_idx % 2 == 1 else "#ffffff"
                    if r_idx == 0:
                        bg = "#5b9bd5"

                    td_style = (
                        f"background-color:{bg}; "
                        "border:1px solid #9cc2e5; vertical-align:top; padding:4px;"
                        "width:263px" if c_idx == 0 else
                        f"background-color:{bg}; border:1px solid #9cc2e5; vertical-align:top; padding:4px; width:303px"
                    )

                    html.append(
                        f'<td style="{td_style}"><p><strong>{text}</strong></p></td>'
                        if c_idx == 0 or r_idx==0 else f'<td style="{td_style}"><p>{text}</p></td>'
                    )
                html.append("</tr>")
            html.append("</tbody></table>")
            return "\n".join(html)
    return ""

# ---------------------------------------Meta Discription---------------------------------------
def extract_meta_description(docx_path):
    doc = Document(docx_path)
    capture = False
    for para in doc.paragraphs:
        text = para.text.strip()
        low = text.lower()
        if not capture and ("introduction" in low):
            capture = True
            continue
        if capture and text:
            return text
    return ""
# --------------------------------------------SEO title-------------------------------------------------------------
def extract_seo_title(docx_path):
    doc = Document(docx_path)
    file_name = os.path.splitext(os.path.basename(docx_path))[0]  # File name without extension
    
    revenue_forecast = ""

    # --- Check tables for Report Coverage ---
    for table in doc.tables:
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]

        if "report attribute" in headers and "details" in headers:
            attr_idx = headers.index("report attribute")
            details_idx = headers.index("details")

            for row in table.rows[1:]:
                attr = row.cells[attr_idx].text.strip().lower()
                details = row.cells[details_idx].text.strip()

                if "revenue forecast in 2030" in attr:
                    # replace USD with $
                    revenue_forecast = details.replace("USD", "$").strip()
                    break

    if revenue_forecast:
        seo_title = f"{file_name} Size ({revenue_forecast}) 2030"
    else:
        seo_title = file_name  # fallback

    return seo_title
# -----------------------------------------------BreadCrumb Text----------------------------------------
def extract_breadcrumb_text(docx_path):
    doc = Document(docx_path)
    file_name = os.path.splitext(os.path.basename(docx_path))[0]  # File name without extension
    
    revenue_forecast = ""

    # --- Check tables for Report Coverage ---
    for table in doc.tables:
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]

        if "report attribute" in headers and "details" in headers:
            attr_idx = headers.index("report attribute")
            details_idx = headers.index("details")

            for row in table.rows[1:]:
                attr = row.cells[attr_idx].text.strip().lower()
                details = row.cells[details_idx].text.strip()

                if "revenue forecast in 2030" in attr:
                    # replace USD with $
                    revenue_forecast = details.replace("USD", "$").strip()
                    break

    if revenue_forecast:
        seo_title = f"{file_name} Report 2030"
    else:
        seo_title = file_name  # fallback

    return seo_title

# ---------------------------------------------SkuCode-Extraction------------------------------
def extract_sku_code(docx_path):
    filename = os.path.basename(docx_path)
    sku_code = os.path.splitext(filename)[0].lower()
    return sku_code
# ---------------------------------------------URLRP------------------------------
def extract_sku_url(docx_path):
    filename = os.path.basename(docx_path)
    sku_code = os.path.splitext(filename)[0].lower()
    return sku_code

# ---------------------------------------------BreadCrumb Schema----------------------------
def extract_breadcrumb_schema(docx_path):
    doc = Document(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    capture = False
    breadcrumb_data = []

    for text in paragraphs:
        low = text.lower()
        if not capture and text.strip().startswith("{"):
            capture = True
        if capture and ("json copy" in low or "faq schema" in low):
            break
        if capture:
            breadcrumb_data.append(text)

    return "".join(breadcrumb_data).strip()
# --------------------------------Schema 2-----------------------
def _get_text(docx_path):
    doc = Document(docx_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip())

def _extract_json_block(text, type_name):
    pat = re.compile(r'"@type"\s*:\s*"' + re.escape(type_name) + r'"')
    m = pat.search(text)
    if not m:
        return ""
    start_idx = text.rfind("{", 0, m.start())
    if start_idx == -1:
        return ""
    depth, i, n = 0, start_idx, len(text)
    block_chars = []
    while i < n:
        ch = text[i]
        block_chars.append(ch)
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                break
        i += 1
    return "".join(block_chars).strip()

def extract_breadcrumb_schema(docx_path):
    text = _get_text(docx_path)
    return _extract_json_block(text, "BreadcrumbList")

def extract_faq_schema(docx_path):
    text = _get_text(docx_path)
    return _extract_json_block(text, "FAQPage")

# ------------------------Methodology-----------------------------------------
import json
import html
def extract_methodology_from_faqschema(docx_path):
    faq_schema_str = extract_faq_schema(docx_path)  
    if not faq_schema_str:
        return ""   
    
    try:
        faq_data = json.loads(faq_schema_str)
    except json.JSONDecodeError:
        return ""   
    
    faqs = []
    q_count = 0
    for item in faq_data.get("mainEntity", []):
        q_count += 1
        question = item.get("name", "").strip()
        answer = item.get("acceptedAnswer", {}).get("text", "").strip()
        if question and answer:
            faqs.append(
                f"<p><strong>Q{q_count}: {html.escape(question)}</strong><br>"
                f"A{q_count}: {html.escape(answer)}</p>"
            )
    
    return "\n".join(faqs)

# ------------------- Merge -------------------
def merge_description_and_coverage(docx_path):
    try:
        desc_html = extract_description(docx_path) or ""
        coverage_html = extract_report_coverage_table_with_style(docx_path) or ""
        merged_html = desc_html + "\n\n" + coverage_html if (desc_html or coverage_html) else ""
        return merged_html
    except Exception as e:
        return f"ERROR: {e}"

# # ------------------- Run Extraction -------------------
folder_path = r"C:\Users\Vishnu\Desktop\oldcontent\23 june\23 june"
output_path = r"C:\Users\Vishnu\Documents\extracted_docs\Extraction_New_Title_Old400.xlsx"

all_data = []

for file in os.listdir(folder_path):
    if not file.endswith(".docx") or file.startswith("~$"):
        continue

    doc_path = os.path.join(folder_path, file)
    print(f"Processing: {file}")
    title = extract_title(doc_path)
    description_html = extract_description(doc_path)
    toc=extract_toc(doc_path)
    methodology=extract_methodology_from_faqschema(doc_path)
    # methodology_html = extract_methodology(doc_path)
    seo_title = extract_seo_title(doc_path)
    breadcrumb_text = extract_breadcrumb_text(doc_path)
    skucode = extract_sku_code(doc_path)
    urlrp = extract_sku_url(doc_path)
    breadcrumb_schema=extract_breadcrumb_schema(doc_path)
    meta=extract_meta_description(doc_path)
    schema2=extract_faq_schema(doc_path)
    report=extract_report_coverage_table_with_style(doc_path)
    merge=merge_description_and_coverage(doc_path)
    chunks = split_into_excel_cells(merge)


    row_data={
        "File": file,
        "Title": title,
        "Description": description_html,
        "TOC":toc,
        "Segmentation":"<p>.</p>",
        "Methodology":methodology,
        "Publish_Date":date.today().strftime("%B %Y"),
        "Currency":"USD",
        "Single Price": 4485,
        "Corporate Price": 6449,
        "skucode": skucode,
        "Total Page":"",
        "Date": date.today().strftime("%Y-%m-%d"),  # always today's date
        "urlNp": urlrp,
        "Meta Discription":meta,
        "Meta Keys":"",
        "Base Year":"2024",
        "history":"2019-2023",
        "Enterprise Price": 8339,
        "SEOTITLE": seo_title,
        "BreadCrumb Text": breadcrumb_text,
        "Schema 1":breadcrumb_schema,
        "Schema 2":schema2,
        "Report":report,
        "Discription":merge
    }
    for i, chunk in enumerate(chunks, start=1):
        row_data[f"Discription_Part{i}"] = chunk
    all_data.append(row_data)

df = pd.DataFrame(all_data)
df.to_excel(output_path, index=False)
print(f"Done! Extracted data saved in {output_path}")
